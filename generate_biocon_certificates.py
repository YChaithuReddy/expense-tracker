"""
Generate editable Biocon Clampon Consolidated Calibration Certificate.
- FluxGen letterhead as header image on every page
- 10 certificate pages, each with its own editable data
- All data editable (serial number, readings, errors, dates)
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


LETTERHEAD = r'C:\Users\chath\Documents\Python code\expense tracker\biocon_images\image1.jpeg'


def set_cell_border(cell, edges=('top', 'bottom', 'left', 'right'), sz='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in edges:
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tcBorders.append(element)
        element.set(qn('w:sz'), sz)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:color'), '000000')


def add_all_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def style_cell(cell, size=9, bold=False, align_center=True):
    if align_center:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for para in cell.paragraphs:
        if align_center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(size)
            if bold:
                run.bold = True


def add_certificate(doc, data, is_first=False):
    """Add one certificate page with editable data."""
    if not is_first:
        # Page break
        doc.add_page_break()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('Calibration Certificate')
    r.bold = True
    r.font.size = Pt(14)
    r.font.underline = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('We certify that')
    r.font.size = Pt(10)

    prod = doc.add_paragraph()
    prod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = prod.add_run('Ultrasonic Flow Meter')
    r.bold = True
    r.font.size = Pt(12)
    r.font.underline = True

    spec = doc.add_paragraph()
    spec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = spec.add_run('With the following specifications')
    r.font.size = Pt(9)
    r.italic = True

    # Specifications table (4 rows x 4 cols)
    spec_table = doc.add_table(rows=4, cols=4)
    spec_table.autofit = False

    # Row 0: Model | Aquagen Clampon USM (merge 2-4)
    r0 = spec_table.rows[0].cells
    r0[0].text = 'Model:'
    r0[1].text = 'Aquagen Clampon USM'
    r0[1].merge(r0[2]).merge(r0[3])

    # Row 1: Sl. No. | serial (merge 2) | Max Pressure | 1.6MPa
    r1 = spec_table.rows[1].cells
    r1[0].text = 'Sl. No.'
    r1[1].text = data['serial']
    r1[2].text = 'Max Pressure'
    r1[3].text = '1.6MPa'

    # Row 2: (blank) | M2 | (blank) | (blank)
    r2 = spec_table.rows[2].cells
    r2[0].text = ''
    r2[1].text = 'M2'
    r2[2].text = ''
    r2[3].text = ''

    # Row 3: Accuracy Class | ±1% | Max Fluid Temp. | 10°C
    r3 = spec_table.rows[3].cells
    r3[0].text = 'Accuracy Class'
    r3[1].text = '\u00b11%'
    r3[2].text = 'Max Fluid Temp.'
    r3[3].text = '10\u00b0C'

    add_all_borders(spec_table)
    for row in spec_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    # Date line
    date_p = doc.add_paragraph()
    r = date_p.add_run('Has been tested and verified on this date: ')
    r.font.size = Pt(9)
    r = date_p.add_run(data['test_date'])
    r.bold = True
    r.font.size = Pt(9)
    r.font.underline = True

    # Flow readings table (5 rows x 11 cols)
    flow_table = doc.add_table(rows=5, cols=11)
    flow_table.autofit = False

    # Header row 1
    hdr1 = flow_table.rows[0].cells
    hdr1[0].text = 'Flow\nvelocity'
    hdr1[1].text = 'Reference Reading (m\u00b3/h)'
    hdr1[1].merge(hdr1[2]).merge(hdr1[3])
    hdr1[4].text = 'Meter Reading (m\u00b3/h)'
    hdr1[4].merge(hdr1[5]).merge(hdr1[6])
    hdr1[7].text = 'Error (%)'
    hdr1[8].text = 'K\u2081'
    hdr1[9].text = 'K\u2082'
    hdr1[10].text = 'K\u2083'

    # Header row 2: numeric sub-labels
    hdr2 = flow_table.rows[1].cells
    hdr2[0].text = ''
    for i, v in enumerate(['1', '2', '3', '1', '2', '3'], start=1):
        hdr2[i].text = v
    hdr2[7].text = ''
    hdr2[8].text = ''
    hdr2[9].text = ''
    hdr2[10].text = 'revised\nerror (%)'

    # Data rows
    for idx, label in enumerate(['Low', 'Middle', 'High']):
        row = flow_table.rows[2 + idx].cells
        vals = [label] + list(data[label.lower()])
        for i, v in enumerate(vals):
            row[i].text = str(v)

    add_all_borders(flow_table)

    for row_idx, row in enumerate(flow_table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(8)
                    if row_idx in (0, 1):
                        run.bold = True

    # Conclusion
    concl = doc.add_paragraph()
    r = concl.add_run('Conclusion:')
    r.bold = True
    r.font.size = Pt(9)
    r = concl.add_run('  It complies with the factory origin standards.  measuring range 100-5000m\u00b3/h.')
    r.font.size = Pt(9)

    # Actual accuracy
    acc_table = doc.add_table(rows=1, cols=2)
    acc_table.rows[0].cells[0].text = 'Actual Accuracy:'
    acc_table.rows[0].cells[1].text = '1'
    add_all_borders(acc_table)
    for row in acc_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    # Operating conditions header
    op_hdr = doc.add_paragraph()
    r = op_hdr.add_run('Operating Conditions at the Flow laboratory')
    r.bold = True
    r.font.size = Pt(9)

    # Op conditions table
    op_table = doc.add_table(rows=2, cols=4)
    op_table.rows[0].cells[0].text = 'Medium'
    op_table.rows[0].cells[1].text = 'Water'
    op_table.rows[0].cells[2].text = 'Medium Temp.'
    op_table.rows[0].cells[3].text = '10\u00b0C'
    op_table.rows[1].cells[0].text = 'Operation Pressure'
    op_table.rows[1].cells[1].text = '1.6MPa'
    op_table.rows[1].cells[2].text = 'Ambient Temp.'
    op_table.rows[1].cells[3].text = '15\u00b0C'
    add_all_borders(op_table)
    for row in op_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    # Calibrating date
    doc.add_paragraph()  # spacer
    cal_p = doc.add_paragraph()
    r = cal_p.add_run('Calibrating date: ')
    r.bold = True
    r.font.size = Pt(9)
    r = cal_p.add_run(data['calibrating_date'])
    r.font.size = Pt(9)
    r.font.underline = True


# === CERTIFICATE DATA (10 certificates) ===
certificates = [
    {
        'serial': '230127-Fn-21948394',
        'test_date': '2024.11.06',
        'calibrating_date': '2024.11.06',
        'low':    ['29.74', '26.14', '34.79', '21.54', '24.89', '30.66', '0.2757', '0.0478', '0.1187', '0.1474'],
        'middle': ['35.99', '48.39', '45.04', '31.80', '48.14', '46.91', '0.1164', '0.0052', '-0.0415', '0.0267'],
        'high':   ['56.33', '52.73', '66.38', '59.14', '55.48', '65.25', '-0.0499', '-0.0522', '0.0170', '-0.0283'],
    },
    {
        'serial': '230127-Fn-21948508',
        'test_date': '2024.11.06',
        'calibrating_date': '2024.11.06',
        'low':    ['23.74', '23.14', '38.79', '25.54', '27.89', '34.66', '-0.0758', '-0.2053', '0.1065', '-0.0582'],
        'middle': ['31.99', '43.39', '40.04', '33.80', '40.14', '47.91', '-0.0566', '0.0749', '-0.1966', '-0.0594'],
        'high':   ['59.33', '57.73', '62.38', '54.14', '52.48', '67.25', '0.0875', '0.0909', '-0.0781', '0.0334'],
    },
    {
        'serial': '230127-Fn-21955169',
        'test_date': '2024.11.06',
        'calibrating_date': '2024.11.06',
        'low':    ['23.74', '25.14', '39.79', '29.54', '27.89', '30.66', '-0.2443', '-0.1094', '0.2295', '-0.0414'],
        'middle': ['37.99', '40.39', '46.04', '36.80', '40.14', '43.91', '0.0313', '0.0062', '0.0463', '0.0279'],
        'high':   ['59.33', '57.73', '68.38', '59.14', '55.48', '65.25', '0.0032', '0.0390', '0.0458', '0.0293'],
    },
    {
        'serial': '230127-Fn-21955158',
        'test_date': '2024.11.06',
        'calibrating_date': '2024.11.06',
        'low':    ['29.74', '27.14', '37.79', '26.54', '23.89', '37.66', '0.1076', '0.1197', '0.0034', '0.0769'],
        'middle': ['39.99', '47.39', '48.04', '38.80', '42.14', '46.91', '0.0298', '0.1108', '0.0235', '0.0547'],
        'high':   ['59.33', '51.73', '61.38', '51.14', '56.48', '68.25', '0.1380', '-0.0918', '-0.1119', '-0.0219'],
    },
    {
        'serial': '230127-Fn-21955121',
        'test_date': '2024.11.06',
        'calibrating_date': '2024.11.06',
        'low':    ['21.74', '26.14', '32.79', '23.54', '21.89', '30.66', '-0.0828', '0.1626', '0.0650', '0.0482'],
        'middle': ['30.99', '46.39', '44.04', '36.80', '45.14', '42.91', '-0.1875', '0.0269', '0.0257', '-0.0450'],
        'high':   ['54.33', '57.73', '63.38', '59.14', '54.48', '63.25', '-0.0885', '0.0563', '0.0021', '-0.0101'],
    },
    {
        'serial': '230127-Fn-21955065',
        'test_date': '2024.11.06',
        'calibrating_date': '2024.11.06',
        'low':    ['29.74', '24.14', '31.79', '23.54', '22.89', '37.66', '0.2085', '0.0518', '-0.1846', '0.0252'],
        'middle': ['36.99', '45.39', '44.04', '39.80', '49.14', '48.91', '-0.0760', '-0.0826', '-0.1106', '-0.0897'],
        'high':   ['56.33', '55.73', '64.38', '59.14', '57.48', '68.25', '-0.0499', '-0.0314', '-0.0601', '-0.0471'],
    },
    {
        'serial': '230127-Fn-21955130',
        'test_date': '2024.11.06',
        'calibrating_date': '2024.11.06',
        'low':    ['25.74', '28.14', '35.79', '24.54', '26.89', '33.66', '0.0467', '-0.0446', '0.0594', '0.0205'],
        'middle': ['38.99', '44.39', '47.04', '37.80', '43.14', '45.91', '0.0305', '0.0282', '0.0240', '0.0276'],
        'high':   ['57.33', '54.73', '65.38', '56.14', '53.48', '64.25', '0.0208', '0.0228', '0.0173', '0.0203'],
    },
    {
        'serial': '230127-Fn-21955114',
        'test_date': '2024.11.06',
        'calibrating_date': '2024.11.06',
        'low':    ['27.74', '25.14', '36.79', '26.54', '24.89', '35.66', '0.0433', '0.0099', '0.0307', '0.0280'],
        'middle': ['34.99', '46.39', '45.04', '35.80', '44.14', '44.91', '-0.0231', '0.0485', '0.0029', '0.0094'],
        'high':   ['55.33', '53.73', '64.38', '57.14', '52.48', '63.25', '-0.0327', '0.0233', '0.0176', '0.0027'],
    },
    {
        'serial': '230127-Fn-21955436',
        'test_date': '2024.11.06',
        'calibrating_date': '2024.11.06',
        'low':    ['24.74', '27.14', '33.79', '25.54', '25.89', '31.66', '-0.0323', '0.0461', '0.0630', '0.0256'],
        'middle': ['36.99', '43.39', '46.04', '38.80', '41.14', '44.91', '-0.0489', '0.0519', '0.0245', '0.0092'],
        'high':   ['58.33', '56.73', '66.38', '57.14', '55.48', '65.25', '0.0204', '0.0220', '0.0170', '0.0198'],
    },
    {
        'serial': '230127-Fn-21955536',
        'test_date': '2024.11.06',
        'calibrating_date': '2024.11.06',
        'low':    ['26.74', '24.14', '37.79', '25.54', '23.89', '36.66', '0.0449', '0.0104', '0.0299', '0.0284'],
        'middle': ['37.99', '45.39', '48.04', '36.80', '46.14', '47.91', '0.0313', '-0.0165', '0.0027', '0.0058'],
        'high':   ['56.33', '55.73', '67.38', '58.14', '54.48', '66.25', '-0.0321', '0.0224', '0.0168', '0.0024'],
    },
]


def add_letterhead_to_header(section):
    """Add the FluxGen letterhead image as a full-page header."""
    header = section.header
    # Remove default paragraph text if any, add image
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(LETTERHEAD, width=Cm(19))
    except Exception as e:
        print(f'Header image failed: {e}')


# === BUILD DOCUMENT ===
doc = Document()

# Set page margins (A4-like, giving space for letterhead)
for section in doc.sections:
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# Page 1: Just the FluxGen letterhead cover page
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run()
run.add_picture(LETTERHEAD, width=Cm(18))

# Add a centered title below the letterhead
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_para.add_run('\n\nCONSOLIDATED CALIBRATION CERTIFICATE\n')
r.bold = True
r.font.size = Pt(18)

sub_title = doc.add_paragraph()
sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_title.add_run('Biocon - Ultrasonic Flow Meters')
r.bold = True
r.font.size = Pt(14)

# Add certificate pages
for idx, cert in enumerate(certificates):
    add_certificate(doc, cert, is_first=False)

# Save
output = r'C:\Users\chath\Documents\Python code\expense tracker\Biocon_Clampon_Calibration_Certificate_v2.docx'
doc.save(output)
print(f'Saved: {output}')
print(f'Pages: 1 cover + {len(certificates)} certificates = {1 + len(certificates)} total')

"""
Generate editable Calibration Certificate Word document (.docx)
Layout matches the provided reference image.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with values like 'single' or {'sz': '4', 'val': 'single', 'color': '000000'}"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            val = kwargs[edge]
            if isinstance(val, str):
                val = {'sz': '4', 'val': val, 'color': '000000'}
            tag = qn(f'w:{edge}')
            element = tcBorders.find(tag)
            if element is None:
                element = OxmlElement(f'w:{edge}')
                tcBorders.append(element)
            for attr, value in val.items():
                element.set(qn(f'w:{attr}'), value)


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_all_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top='single', bottom='single', left='single', right='single')


def run(text, bold=False, size=10, underline=False):
    return {'text': text, 'bold': bold, 'size': size, 'underline': underline}


doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# === TITLE ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Calibration Certificate')
r.bold = True
r.font.size = Pt(16)
r.font.underline = True

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('We certify that')
r.font.size = Pt(11)

# Product name
prod = doc.add_paragraph()
prod.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = prod.add_run('Ultrasonic Flow Meter')
r.bold = True
r.font.size = Pt(13)
r.font.underline = True

spec = doc.add_paragraph()
spec.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = spec.add_run('With the following specifications')
r.font.size = Pt(10)
r.italic = True

# === SPECIFICATIONS TABLE ===
spec_table = doc.add_table(rows=3, cols=4)
spec_table.autofit = False

# Row 1: Model | Aquagen Clampon USM (spans 3)
row1 = spec_table.rows[0].cells
row1[0].text = 'Model:'
row1[1].text = 'Aquagen Clampon USM'
row1[1].merge(row1[2]).merge(row1[3])

# Row 2: Sl. No. | 230127-Fn-21956536 (spans) | Max Pressure | 1.6MPa
row2 = spec_table.rows[1].cells
row2[0].text = 'Sl. No.'
row2[1].text = '230127-Fn-21956536'
row2[2].text = 'Max Pressure'
row2[3].text = '1.6MPa'

# Row 3 (added): M2 | (blank) | (blank) | (blank) -- actually from image it's merged differently
# Actually from image: 3rd row has "M2" under Sl. No. column spanning with 1.6MPa cell on right
# Let's simplify: add another row for M2
spec_table.add_row()
row3 = spec_table.rows[2].cells
row3[0].text = ''
row3[1].text = 'M2'
row3[2].text = ''
row3[3].text = ''

# Row 4: Accuracy Class | ±1% | Max Fluid Temp. | 10°C
row4 = spec_table.rows[3].cells
row4[0].text = 'Accuracy Class'
row4[1].text = '±1%'
row4[2].text = 'Max Fluid Temp.'
row4[3].text = '10°C'

add_all_borders(spec_table)

# Style the spec table
for row in spec_table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

# === DATE LINE ===
date_p = doc.add_paragraph()
r = date_p.add_run('Has been tested and verified on this date: ')
r.font.size = Pt(10)
r = date_p.add_run('2024.11.06')
r.bold = True
r.font.size = Pt(10)
r.font.underline = True

# === FLOW READINGS TABLE ===
# Header row 1: Flow velocity | Reference Reading (m³/h) (spans 3) | Meter Reading (m³/h) (spans 3) | Error (%) | K1 | K2 | K3 | revised error
# Header row 2: (blank) | 1 | 2 | 3 | 1 | 2 | 3 | (blank) | (blank) | (blank) | (blank) | (%)

flow_table = doc.add_table(rows=5, cols=11)
flow_table.autofit = False

# Header row 1 - with merges
hdr1 = flow_table.rows[0].cells
hdr1[0].text = 'Flow\nvelocity'
# Reference Reading spans 1-3
hdr1[1].text = 'Reference Reading (m\u00b3/h)'
hdr1[1].merge(hdr1[2]).merge(hdr1[3])
# Meter Reading spans 4-6
hdr1[4].text = 'Meter Reading (m\u00b3/h)'
hdr1[4].merge(hdr1[5]).merge(hdr1[6])
hdr1[7].text = 'Error (%)'
hdr1[8].text = 'K\u2081'
hdr1[9].text = 'K\u2082'
hdr1[10].text = 'K\u2083'

# Row 2: sub-headers for 1, 2, 3
hdr2 = flow_table.rows[1].cells
hdr2[0].text = ''
hdr2[1].text = '1'
hdr2[2].text = '2'
hdr2[3].text = '3'
hdr2[4].text = '1'
hdr2[5].text = '2'
hdr2[6].text = '3'
hdr2[7].text = ''
hdr2[8].text = ''
hdr2[9].text = ''
hdr2[10].text = 'revised error\n(%)'

# Data rows
def fill_row(cells, values):
    for i, v in enumerate(values):
        cells[i].text = str(v)

fill_row(flow_table.rows[2].cells, ['Low', '23.74', '26.14', '34.79', '21.54', '23.89', '32.66', '0.0927', '0.0861', '0.0612', '0.0800'])
fill_row(flow_table.rows[3].cells, ['Middle', '36.99', '49.39', '49.04', '39.80', '43.14', '40.91', '-0.0760', '0.1265', '0.1658', '0.0721'])
fill_row(flow_table.rows[4].cells, ['High', '52.33', '51.73', '63.38', '53.14', '51.48', '66.25', '-0.0155', '0.0048', '-0.0453', '-0.0186'])

add_all_borders(flow_table)

# Center-align all cells in flow table, set font size
for row in flow_table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)

# Bold the header rows
for i in [0, 1]:
    for cell in flow_table.rows[i].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

# === CONCLUSION ===
concl = doc.add_paragraph()
r = concl.add_run('Conclusion:')
r.bold = True
r.font.size = Pt(10)
r = concl.add_run(' It complies with the factory origin standards.  measuring range 100-5000m\u00b3/h.')
r.font.size = Pt(10)

# === ACTUAL ACCURACY TABLE ===
acc_table = doc.add_table(rows=1, cols=2)
acc_table.rows[0].cells[0].text = 'Actual Accuracy:'
acc_table.rows[0].cells[1].text = '1'
add_all_borders(acc_table)
for row in acc_table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

# === OPERATING CONDITIONS HEADER ===
op_hdr = doc.add_paragraph()
r = op_hdr.add_run('Operating Conditions at the Flow laboratory')
r.bold = True
r.font.size = Pt(10)

# === OPERATING CONDITIONS TABLE ===
op_table = doc.add_table(rows=2, cols=4)
op_rows = op_table.rows

op_rows[0].cells[0].text = 'Medium'
op_rows[0].cells[1].text = 'Water'
op_rows[0].cells[2].text = 'Medium Temp.'
op_rows[0].cells[3].text = '10\u00b0C'

op_rows[1].cells[0].text = 'Operation Pressure'
op_rows[1].cells[1].text = '1.6MPa'
op_rows[1].cells[2].text = 'Ambient Temp.'
op_rows[1].cells[3].text = '15\u00b0C'

add_all_borders(op_table)
for row in op_table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

# === LOCATION FIELD (editable) ===
loc_p = doc.add_paragraph()
r = loc_p.add_run('Location: ')
r.bold = True
r.font.size = Pt(10)
r = loc_p.add_run('________________________________')
r.font.size = Pt(10)

# === CALIBRATING DATE (bottom-left, editable) ===
doc.add_paragraph()  # spacer
cal_p = doc.add_paragraph()
r = cal_p.add_run('Calibrating date: ')
r.bold = True
r.font.size = Pt(10)
r = cal_p.add_run('2024.11.06')
r.font.size = Pt(10)
r.font.underline = True

# === SAVE ===
output = r'C:\Users\chath\Documents\Python code\expense tracker\Calibration_Certificate.docx'
doc.save(output)
print(f'Saved to: {output}')

"""Read the Biocon certificate to understand its structure."""
from docx import Document

path = r'C:\Users\chath\Downloads\Biocon Clampon Consolidated Calibration Certificate.docx'
doc = Document(path)

print('=' * 70)
print('PARAGRAPHS')
print('=' * 70)
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        style = para.style.name if para.style else 'Normal'
        print(f'[{i}] ({style}) {text[:200]}')

print()
print('=' * 70)
print(f'TABLES: {len(doc.tables)}')
print('=' * 70)
for ti, table in enumerate(doc.tables):
    print(f'\n--- Table {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---')
    for ri, row in enumerate(table.rows):
        cells_text = [cell.text.strip()[:60] for cell in row.cells]
        print(f'  Row {ri}: {cells_text}')

print()
print('=' * 70)
print('SECTIONS / PAGE BREAKS')
print('=' * 70)
print(f'Sections: {len(doc.sections)}')
for si, section in enumerate(doc.sections):
    print(f'  Section {si}: top={section.top_margin}, left={section.left_margin}, page_w={section.page_width}')

# Check for images
print()
print('=' * 70)
print('INLINE SHAPES / IMAGES')
print('=' * 70)
print(f'Inline shapes: {len(doc.inline_shapes)}')
